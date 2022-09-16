# -*-coding: utf-8 -*-
import os
from docx import Document   #用来建立一个word对象
import shutil
import win32com.client as win32
doc = Document()
xls_app = win32.gencache.EnsureDispatch('Excel.Application')
wb = xls_app.Workbooks.Add()
ws = wb.Worksheets(1)
ws.Name = '底稿目录'
xls_app.Visible = True

startpath = r"D:\\SynologyDrive\\FY01a【归档】物业尾款一期\\2 尽职调查阶段\\"
for root, dirs, files in os.walk(startpath):
    level = root.replace(startpath, '').count(os.sep)
    indent = ' ' * 4 * (level)
    para_heading = doc.add_paragraph('{}{}/'.format(indent,os.path.basename(root)))
    print("'{}/".format(os.path.basename(root)))
    subindent = ' ' * 4 * (level + 1)
    for f in files:
        para_heading2 = doc.add_paragraph('{}{}'.format(subindent, f))
        # print('{}{}'.format(subindent, f))

doc.save(startpath+"目录【终稿】.docx")
